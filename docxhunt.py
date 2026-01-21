#!/usr/bin/env python3
import argparse
import requests
from docx import Document
import re
import os
import sys
import math
import unicodedata

# ================= CONFIG =================
DOWNLOAD_DIR = "docx_files"
MAX_FILE_SIZE = 200 * 1024 * 1024  # 200 MB
os.makedirs(DOWNLOAD_DIR, exist_ok=True)

HEADERS = {"User-Agent": "DOCX-Sensitive-Scanner/7.0"}

CONTEXT_KEYWORDS = [
    "password", "passwd", "pwd", "secret", "token", "api",
    "key", "credential", "auth", "login", "confidential", "private"
]

PLACEHOLDERS = {"xxxxxxxx", "1234567890", "0000000000", "abcdef"}

SKIP_LOG_FILE = "skipped_files.txt"

# ================= BUILT-IN REGEX =================
BUILTIN_PATTERNS = {
    "EMAIL": r"[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}",
    "PHONE": r"\+?\d{1,3}[\s.-]?\(?\d{2,4}\)?[\s.-]?\d{3,4}[\s.-]?\d{4}",
    "AADHAAR": r"\b[2-9][0-9]{3}\s?[0-9]{4}\s?[0-9]{4}\b",
    "PAN": r"\b[A-Z]{5}[0-9]{4}[A-Z]\b",
    "SSN": r"\b\d{3}-\d{2}-\d{4}\b",
    "CREDIT_CARD": r"\b(?:4[0-9]{12}(?:[0-9]{3})?|5[1-5][0-9]{14}|3[47][0-9]{13})\b",
    "IBAN": r"\b[A-Z]{2}[0-9]{2}[A-Z0-9]{11,30}\b",
    "PASSWORD": r"(?i)\b(password|passwd|pwd)\b\s*[:=]\s*['\"].{4,100}['\"]",
    "JWT": r"eyJ[A-Za-z0-9_-]+\.eyJ[A-Za-z0-9_-]+\.?[A-Za-z0-9_-]*",
    "BEARER_TOKEN": r"(?i)bearer\s+[a-z0-9\-._~+/]+=*",
    "AWS_ACCESS_KEY": r"AKIA[0-9A-Z]{16}",
    "AWS_SECRET_KEY": r"(?i)aws(.{0,20})?(secret|private).{0,20}['\"][A-Za-z0-9/+=]{40}['\"]",
    "GOOGLE_API_KEY": r"AIza[0-9A-Za-z\-_]{35}",
    "AZURE_SECRET": r"(?i)azure(.{0,20})?(key|secret).{0,20}['\"][A-Za-z0-9/+=]{40,}['\"]",
    "DB_CONN": r"(?i)(mysql|postgres|mongodb|redis|sqlserver):\/\/[^'\"\s]+",
    "INTERNAL_IP": r"\b(10\.|192\.168\.|172\.(1[6-9]|2[0-9]|3[0-1]))\d+\.\d+\b",
    "SSH_PRIVATE_KEY": r"-----BEGIN (RSA|DSA|EC|OPENSSH) PRIVATE KEY-----",
    "HIGH_ENTROPY": r"[A-Za-z0-9+/=]{32,}",
}

# ================= SEVERITY =================
BASE_SCORE = {
    "EMAIL": 10, "PHONE": 10, "AADHAAR": 30, "PAN": 30, "SSN": 40,
    "CREDIT_CARD": 40, "IBAN": 35,
    "PASSWORD": 60, "JWT": 90, "BEARER_TOKEN": 80,
    "AWS_ACCESS_KEY": 95, "AWS_SECRET_KEY": 95,
    "GOOGLE_API_KEY": 90, "AZURE_SECRET": 90,
    "DB_CONN": 70, "INTERNAL_IP": 20,
    "SSH_PRIVATE_KEY": 100, "HIGH_ENTROPY": 30,
}

# ================= UTILS =================
def entropy(s):
    probs = [s.count(c) / len(s) for c in set(s)]
    return -sum(p * math.log2(p) for p in probs)

def classify(score):
    if score >= 90:
        return "CRITICAL"
    if score >= 70:
        return "HIGH"
    if score >= 40:
        return "MED"
    return "LOW"

def reduce_fp(val):
    v = val.lower()
    return v in PLACEHOLDERS or len(set(v)) == 1

def context_boost(text, val):
    idx = text.lower().find(val.lower())
    if idx == -1:
        return 0
    window = text[max(0, idx - 60): idx + 60].lower()
    return 20 if any(k in window for k in CONTEXT_KEYWORDS) else 0

def is_template(text):
    return bool(re.search(r"_{3,}|XXXXX|SAMPLE|TEMPLATE", text, re.I))

# ================= SCAN DOCX =================
def scan_docx(path, compiled_patterns, log):
    try:
        doc = Document(path)
    except Exception as e:
        log(f"[v] DOCX parse error: {e}")
        return None, None

    text = "\n".join(p.text for p in doc.paragraphs)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                text += "\n" + cell.text

    for section in doc.sections:
        for p in section.header.paragraphs:
            text += "\n" + p.text
        for p in section.footer.paragraphs:
            text += "\n" + p.text

    text = unicodedata.normalize("NFKC", text)
    template = is_template(text)

    results, seen = [], set()

    for label, regex in compiled_patterns.items():
        for m in regex.findall(text):
            val = m if isinstance(m, str) else m[0]
            key = f"{label}:{val}"

            if key in seen or reduce_fp(val):
                continue
            seen.add(key)

            score = BASE_SCORE.get(label, 30)
            score += context_boost(text, val)

            if label == "HIGH_ENTROPY":
                if entropy(val) < 4.2 or context_boost(text, val) == 0:
                    continue
                score += 20

            if template:
                score = max(score - 30, 10)

            results.append({
                "type": label,
                "value": val[:80],
                "severity": classify(score)
            })

    results.sort(key=lambda x: ["LOW", "MED", "HIGH", "CRITICAL"].index(x["severity"]))
    return results, template

# ================= MAIN =================
def main():
    parser = argparse.ArgumentParser(
        description="DOCX Sensitive Data Scanner (Production Ready)",
        formatter_class=argparse.RawTextHelpFormatter
    )

    parser.add_argument("-i", "--input", required=True, help="Input file with DOCX URLs")
    parser.add_argument("-o", "--output", help="Write results to output file")
    parser.add_argument("-r", "--regex", help="Custom regex file")
    parser.add_argument("-v", "--verbose", action="store_true", help="Verbose output")
    parser.add_argument("--skip-show", action="store_true", help="Save skipped files with reasons")

    args = parser.parse_args()

    def log(msg):
        if args.verbose:
            print(msg)

    total_urls = scanned = skipped = failed = 0
    skipped_not_docx = skipped_too_large = skipped_parse_error = 0

    skip_records = []

    patterns = dict(BUILTIN_PATTERNS)
    if args.regex:
        patterns.update({})

    compiled_patterns = {k: re.compile(v) for k, v in patterns.items()}

    with open(args.input) as f:
        urls = [u.strip() for u in f if u.strip()]

    output_lines = []

    for url in urls:
        total_urls += 1
        output_lines.append(f"[*] {url}")

        if not url.lower().endswith(".docx"):
            skipped += 1
            skipped_not_docx += 1
            reason = "Not a DOCX file"
            output_lines.append("    [!] Skipped (not DOCX)\n")
            skip_records.append(f"{url} | {reason}")
            continue

        fname = os.path.join(DOWNLOAD_DIR, os.path.basename(url))
        log(f"[v] Downloading {fname}")

        try:
            r = requests.get(url, headers=HEADERS, timeout=30)
            r.raise_for_status()
            with open(fname, "wb") as f:
                f.write(r.content)
        except Exception:
            failed += 1
            reason = "Download failed"
            output_lines.append("    [!] Download failed\n")
            skip_records.append(f"{url} | {reason}")
            continue

        if os.path.getsize(fname) > MAX_FILE_SIZE:
            skipped += 1
            skipped_too_large += 1
            reason = "File too large (>200MB)"
            output_lines.append("    [!] Skipped (file too large)\n")
            skip_records.append(f"{url} | {reason}")
            continue

        result = scan_docx(fname, compiled_patterns, log)
        if result == (None, None):
            skipped += 1
            skipped_parse_error += 1
            reason = "DOCX parse error"
            output_lines.append("    [!] Skipped (DOCX parse error)\n")
            skip_records.append(f"{url} | {reason}")
            continue

        scanned += 1
        findings, template = result

        if not findings:
            output_lines.append("    [+] No sensitive data\n")
            continue

        output_lines.append(f"    [!] Findings (Template={template}):")
        for fnd in findings:
            output_lines.append(f"        [{fnd['severity']}] {fnd['type']} → {fnd['value']}")
        output_lines.append("")

    result = "\n".join(output_lines)
    print(result)

    summary = (
        "\n[+] Scan completed\n"
        f"    Total URLs        : {total_urls}\n"
        f"    Scanned DOCX      : {scanned}\n"
        f"    Skipped files     : {skipped}\n"
        f"        - Not DOCX    : {skipped_not_docx}\n"
        f"        - Too large   : {skipped_too_large}\n"
        f"        - Parse error : {skipped_parse_error}\n"
        f"    Failed download   : {failed}\n"
    )

    print(summary)

    if args.skip_show and skip_records:
        with open(SKIP_LOG_FILE, "w") as f:
            for line in skip_records:
                f.write(line + "\n")
        print(f"[+] Skipped files list saved to {SKIP_LOG_FILE}")

    if args.output:
        with open(args.output, "w") as f:
            f.write(result + summary)

if __name__ == "__main__":
    main()
